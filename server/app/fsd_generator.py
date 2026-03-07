from __future__ import annotations

import json
import logging
import re

from docx import Document

from .fsd_template import FSD_SECTIONS, FSD_STRUCTURE, build_fsd_prompt
from .llm_service import generate_text

logger = logging.getLogger(__name__)

_FUNCTIONAL_SECTION_KEYS = [
    "Functional Specification - General Requirements",
    "Functional Specification - Visual Requirements",
    "Functional Specification - Error Handling",
]


def _normalize_functionality_steps(text: str) -> str:
    cleaned = re.sub(r"\s+", " ", str(text or "")).strip()
    if not cleaned:
        return "1. TBD"
    numbered_steps = [
        re.sub(r"^\d+\.\s*", "", part).strip(" ;")
        for part in re.findall(r"\d+\.\s*[^\d]+?(?=(?:\s+\d+\.\s)|$)", cleaned)
    ]
    steps = [step for step in numbered_steps if step]
    if not steps:
        sentence_steps = [segment.strip(" .;") for segment in re.split(r"[.;]\s+", cleaned) if segment.strip()]
        steps = sentence_steps or [cleaned]
    return " ".join(f"{idx}. {step}" for idx, step in enumerate(steps, start=1))


def _parse_functional_row(item: str) -> tuple[str, str, str, str]:
    text = str(item or "").strip()
    pattern = re.compile(
        r"Viewport:\s*(.*?)\s*\|\|\s*"
        r"Visual Reference:\s*(.*?)\s*\|\|\s*"
        r"Element:\s*(.*?)\s*\|\|\s*"
        r"Element Functionality:\s*(.*)$",
        re.IGNORECASE,
    )
    match = pattern.match(text)
    if match:
        viewport, visual_ref, element, functionality = [part.strip() for part in match.groups()]
    else:
        viewport = "All Viewports"
        visual_ref = "TBD screenshot"
        element = text.split(":", 1)[0].strip() if ":" in text else text
        element = element or "TBD element"
        functionality = text

    viewport = viewport or "All Viewports"
    visual_ref = visual_ref or "TBD screenshot"
    element = element or "TBD element"
    normalized_functionality = _normalize_functionality_steps(functionality)
    return viewport, visual_ref, element, normalized_functionality


def _build_functional_rows(fsd_json: dict) -> list[tuple[str, str, str, str]]:
    rows: list[tuple[str, str, str, str]] = []
    for key in _FUNCTIONAL_SECTION_KEYS:
        for item in fsd_json.get(key, []) or []:
            rows.append(_parse_functional_row(item))
    if rows:
        return rows
    return [(
        "All Viewports",
        "TBD screenshot",
        "No functional topics captured",
        "1. Add functional topics to generate sequential flow",
    )]


def _empty_fsd() -> dict:
    return {section: [] for section in FSD_SECTIONS}


def _parse_fsd_json(text: str) -> dict:
    start = text.find("{")
    end = text.rfind("}")
    if start == -1 or end == -1 or end <= start:
        raise ValueError("No JSON object found")
    payload = json.loads(text[start : end + 1])
    if not isinstance(payload, dict):
        raise ValueError("FSD JSON must be an object")
    for section in FSD_SECTIONS:
        payload.setdefault(section, [])
        if not isinstance(payload[section], list):
            payload[section] = [str(payload[section])]
        payload[section] = [str(item) for item in payload[section]]
    return payload


def generate_fsd_json(gap_results: list[dict]) -> dict:
    prompt = build_fsd_prompt(gap_results)
    response_text = generate_text(prompt)
    try:
        return _parse_fsd_json(response_text)
    except Exception as exc:
        logger.warning("FSD JSON parse failed: %s", exc)
        return _empty_fsd()


def render_fsd_text(fsd_json: dict) -> str:
    lines: list[str] = [
        "# Table of Contents",
        "- [Overview](#overview)",
        "- [Background](#background)",
        "  - [Scope](#scope)",
        "  - [Out of Scope](#out-of-scope)",
        "- [Functional Specification](#functional-specification)",
        "  - [Sequential Flow and Element Details](#sequential-flow-and-element-details)",
        "",
        "# Overview",
    ]

    overview_items = fsd_json.get("Overview", [])
    if not overview_items:
        lines.append("- No items.")
    else:
        for item in overview_items:
            lines.append(f"- {item}")

    lines.extend(["", "# Background", "## Scope"])
    scope_items = fsd_json.get("Background - Scope", [])
    if not scope_items:
        lines.append("- No items.")
    else:
        for item in scope_items:
            lines.append(f"- {item}")

    lines.extend(["", "## Out of Scope"])
    out_scope_items = fsd_json.get("Background - Out of Scope", [])
    if not out_scope_items:
        lines.append("- No items.")
    else:
        for item in out_scope_items:
            lines.append(f"- {item}")

    lines.extend(["", "# Functional Specification", "## Sequential Flow and Element Details"])
    lines.append("| Sl No | Viewport | Visual Reference | Element | Element Functionality |")
    lines.append("| --- | --- | --- | --- | --- |")
    for index, row in enumerate(_build_functional_rows(fsd_json), start=1):
        viewport, visual_ref, element, functionality = [cell.replace("|", "/") for cell in row]
        lines.append(f"| {index} | {viewport} | {visual_ref} | {element} | {functionality} |")

    return "\n".join(lines)


def generate_fsd(gap_results: list[dict]) -> str:
    fsd_json = generate_fsd_json(gap_results)
    return render_fsd_text(fsd_json)


def generate_fsd_docx(fsd_json: dict) -> Document:
    doc = Document()
    doc.add_heading("Functional Specification Document", level=1)
    if not fsd_json:
        doc.add_paragraph("No content generated.")
        return doc

    last_parent = None
    for parent, child, key in FSD_STRUCTURE:
        if parent != last_parent:
            doc.add_heading(parent, level=2)
            last_parent = parent
        if child:
            doc.add_heading(child, level=3)
        items = fsd_json.get(key, [])
        if not items:
            doc.add_paragraph("No items.")
            continue
        for item in items:
            doc.add_paragraph(str(item), style="List Bullet")
    return doc


def generate_fsd_docx_from_text(fsd_text: str) -> Document:
    doc = Document()
    doc.add_heading("Functional Specification Document", level=1)
    content = fsd_text.strip()
    if not content:
        doc.add_paragraph("No content generated.")
        return doc

    for raw_line in content.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("#"):
            heading_level = min(3, len(line) - len(line.lstrip("#")))
            heading_text = line.lstrip("#").strip()
            if heading_text:
                doc.add_heading(heading_text, level=heading_level + 1)
            continue
        if re.match(r"^[-*]\s+", line):
            doc.add_paragraph(re.sub(r"^[-*]\s+", "", line), style="List Bullet")
            continue
        if re.match(r"^\d+\.\s+", line):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", line), style="List Number")
            continue
        if line.endswith(":") and len(line) < 140:
            doc.add_heading(line[:-1].strip(), level=2)
            continue
        doc.add_paragraph(line)
    return doc


def export_fsd_to_docx(fsd_json: dict) -> Document:
    return generate_fsd_docx(fsd_json)
